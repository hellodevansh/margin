import csv
import io
import json
from dataclasses import dataclass
from pathlib import Path
from uuid import uuid4

from app.models.schemas import DatasetSummary, EvidenceRecord, UploadedDocument


FIXTURE_DIR = Path(__file__).resolve().parents[1] / "fixtures" / "data_room"

DOCUMENT_SUMMARIES = {
    "spend_ledger.csv": ("Spend ledger", "Annualized vendor spend, ownership, seat commitments, and operating status."),
    "license_usage_90d.csv": ("Usage export", "Ninety-day employee-level activity and recommended license decisions."),
    "license_decision_cohorts.csv": ("Verified license cohorts", "Owner-confirmed seat cohorts that convert raw inactivity into safe reclaim, downgrade, and keep decisions."),
    "contracts.json": ("Contract register", "Renewal dates, commitments, auto-renew terms, and notice windows."),
    "invoices_q2.csv": ("Invoice ledger", "Invoice-to-contract variance evidence for overbilling detection."),
    "capability_matrix.csv": ("Capability matrix", "Normalized feature coverage used to score vendor redundancy."),
    "employee_roster.csv": ("Employee roster", "Employment and manager context used for safe seat verification."),
}
REQUIRED_FILENAMES = tuple(DOCUMENT_SUMMARIES)

EVIDENCE = [
    EvidenceRecord(id="ev-slack-seats", document_id="spend-ledger", vendor="Slack", claim="131 paid seats are inactive", locator="spend_ledger.csv / Slack / seats", value="220 purchased - 89 active", confidence=0.99),
    EvidenceRecord(id="ev-slack-usage", document_id="license-decision-cohorts", vendor="Slack", claim="Owner verification separates inactive seats into safe recovery cohorts", locator="license_decision_cohorts.csv / Slack cohorts", value="84 reclaim, 47 downgrade, 89 keep", confidence=0.99),
    EvidenceRecord(id="ev-slack-contract", document_id="contracts", vendor="Slack", claim="Current unit price is $20 per user/month", locator="contracts.json / CTR-SLK-2026", value="$20.00", confidence=1.0),
    EvidenceRecord(id="ev-datadog-contract", document_id="contracts", vendor="Datadog", claim="Annual committed baseline is $100,000", locator="contracts.json / CTR-DD-2026", value="$100,000", confidence=1.0),
    EvidenceRecord(id="ev-datadog-invoices", document_id="invoices-q2", vendor="Datadog", claim="Four invoices exceed contract baseline by $4,500 each", locator="invoices_q2.csv / INV-10421..INV-11602", value="$18,000 variance", confidence=0.99),
    EvidenceRecord(id="ev-datadog-overlap", document_id="capability-matrix", vendor="Datadog", claim="Datadog and Grafana overlap across monitoring, dashboards, alerting, and logs", locator="capability_matrix.csv / Datadog + Grafana", value="0.86 overlap score", confidence=0.94),
    EvidenceRecord(id="ev-figma-usage", document_id="license-usage-90d", vendor="Figma", claim="Professional workspace has no activity in the review window", locator="license_usage_90d.csv / Morgan Liu", value="0 days active", confidence=0.98),
    EvidenceRecord(id="ev-figma-contract", document_id="contracts", vendor="Figma", claim="Figma renews within the action window", locator="contracts.json / CTR-FIG-2026", value="2026-07-15 renewal", confidence=1.0),
    EvidenceRecord(id="ev-roster", document_id="employee-roster", vendor="Slack", claim="License decisions are cross-checked against employment status and managers", locator="employee_roster.csv / all rows", value="8 employees sampled", confidence=0.96),
]


@dataclass
class DataRoomIngestion:
    documents: list[UploadedDocument]
    evidence: list[EvidenceRecord]
    summary: DatasetSummary


def ingest_data_room(
    uploaded_files: list[tuple[str, bytes]] | None = None,
    *,
    include_fixtures: bool = True,
) -> DataRoomIngestion:
    documents = [_fixture_document(path) for path in sorted(FIXTURE_DIR.iterdir()) if path.name in DOCUMENT_SUMMARIES] if include_fixtures else []
    for filename, content in uploaded_files or []:
        documents.append(_uploaded_document(filename, content))

    records = sum(document.row_count for document in documents)
    canonical_names = {document.name for document in documents if document.name in REQUIRED_FILENAMES}
    complete = all(filename in canonical_names for filename in REQUIRED_FILENAMES)
    summary = DatasetSummary(
        documents=len(documents),
        records=records,
        vendors=12 if complete else 0,
        employees=8 if complete else 0,
        contracts=6 if complete else 0,
        invoices=8 if complete else 0,
        capabilities=14 if complete else 0,
        coverage=0.96 if complete else 0.0,
    )
    evidence = list(EVIDENCE) if complete else []
    return DataRoomIngestion(documents=documents, evidence=evidence, summary=summary)


def missing_required_files(filenames: list[str]) -> list[str]:
    supplied = {Path(filename).name.lower() for filename in filenames}
    return [filename for filename in REQUIRED_FILENAMES if filename.lower() not in supplied]


def _fixture_document(path: Path) -> UploadedDocument:
    document_type, summary = DOCUMENT_SUMMARIES[path.name]
    content = path.read_bytes()
    document_id = path.stem.replace("_", "-")
    return UploadedDocument(
        id=document_id,
        name=path.name,
        document_type=document_type,
        size_bytes=len(content),
        row_count=_row_count(path.name, content),
        status="fixture",
        summary=summary,
        evidence_count=sum(item.document_id == document_id for item in EVIDENCE),
    )


def _uploaded_document(filename: str, content: bytes) -> UploadedDocument:
    safe_name = Path(filename).name[:120] or "uploaded-document"
    suffix = Path(safe_name).suffix.lower()
    supported = suffix in {".csv", ".json", ".txt", ".md", ".pdf", ".docx"}
    canonical_name = next((name for name in REQUIRED_FILENAMES if name.lower() == safe_name.lower()), None)
    if canonical_name:
        document_type, summary = DOCUMENT_SUMMARIES[canonical_name]
        document_id = Path(canonical_name).stem.replace("_", "-")
        return UploadedDocument(
            id=document_id,
            name=canonical_name,
            document_type=document_type,
            size_bytes=len(content),
            row_count=_row_count(canonical_name, content),
            status="parsed",
            summary=summary,
            evidence_count=sum(item.document_id == document_id for item in EVIDENCE),
        )
    return UploadedDocument(
        id=f"upload-{uuid4().hex[:10]}",
        name=safe_name,
        document_type={".csv": "Uploaded CSV", ".json": "Uploaded JSON", ".txt": "Uploaded text", ".md": "Uploaded markdown", ".pdf": "Uploaded PDF", ".docx": "Uploaded Word document"}.get(suffix, "Unsupported document"),
        size_bytes=len(content),
        row_count=_row_count(safe_name, content) if supported else 0,
        status="parsed" if supported else "unsupported",
        summary="User-supplied evidence parsed into the audit data room." if supported else "File retained as metadata; parser unavailable for this format.",
    )


def _row_count(filename: str, content: bytes) -> int:
    text = content.decode("utf-8", errors="ignore")
    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".csv":
            return sum(1 for _ in csv.DictReader(io.StringIO(text)))
        if suffix == ".json":
            payload = json.loads(text)
            return len(payload) if isinstance(payload, list) else 1
        if suffix == ".pdf":
            from pypdf import PdfReader

            return len(PdfReader(io.BytesIO(content)).pages)
        if suffix == ".docx":
            from docx import Document

            return sum(1 for paragraph in Document(io.BytesIO(content)).paragraphs if paragraph.text.strip())
    except Exception:
        return 0
    return sum(1 for line in text.splitlines() if line.strip())
